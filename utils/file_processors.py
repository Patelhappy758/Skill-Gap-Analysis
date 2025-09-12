"""
File processing utilities for document parsing and text extraction.
Supports PDF, DOCX, and TXT file formats.
"""

import os
import re
from typing import Optional, Dict, Any
import PyPDF2
import docx
from io import BytesIO


class DocumentProcessor:
    """Handles document parsing and text extraction for various file formats."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text(self, file_path: str = None, file_content: bytes = None, 
                    file_type: str = None) -> Dict[str, Any]:
        """
        Extract text from various document formats.
        
        Args:
            file_path (str): Path to the file
            file_content (bytes): File content as bytes (for uploaded files)
            file_type (str): File extension (.pdf, .docx, .txt)
        
        Returns:
            Dict containing extracted text, metadata, and success status
        """
        result = {
            'text': '',
            'success': False,
            'error': None,
            'metadata': {}
        }
        
        try:
            if file_path:
                file_type = os.path.splitext(file_path)[1].lower()
                
            if file_type == '.pdf':
                result = self._extract_pdf(file_path, file_content)
            elif file_type == '.docx':
                result = self._extract_docx(file_path, file_content)
            elif file_type == '.txt':
                result = self._extract_txt(file_path, file_content)
            else:
                result['error'] = f"Unsupported file format: {file_type}"
                
        except Exception as e:
            result['error'] = str(e)
            
        return result
    
    def _extract_pdf(self, file_path: str = None, file_content: bytes = None) -> Dict[str, Any]:
        """Extract text from PDF files."""
        result = {'text': '', 'success': False, 'error': None, 'metadata': {}}
        
        try:
            if file_content:
                # Handle uploaded file content
                pdf_file = BytesIO(file_content)
                reader = PyPDF2.PdfReader(pdf_file)
            else:
                # Handle file path
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    page_count = len(reader.pages)
                    
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                    
                    result.update({
                        'text': text.strip(),
                        'success': True,
                        'metadata': {
                            'page_count': page_count,
                            'file_type': 'PDF'
                        }
                    })
                    return result
            
            # For uploaded content (outside the file context)
            text = ""
            page_count = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text += page_text + "\n"
            
            result.update({
                'text': text.strip(),
                'success': True,
                'metadata': {
                    'page_count': page_count,
                    'file_type': 'PDF'
                }
            })
            
        except Exception as e:
            result['error'] = f"Error processing PDF: {str(e)}"
            
        return result
    
    def _extract_docx(self, file_path: str = None, file_content: bytes = None) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        result = {'text': '', 'success': False, 'error': None, 'metadata': {}}
        
        try:
            if file_content:
                # Handle uploaded file content
                docx_file = BytesIO(file_content)
                doc = docx.Document(docx_file)
            else:
                # Handle file path
                doc = docx.Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            text = "\n".join(paragraphs)
            
            result.update({
                'text': text,
                'success': True,
                'metadata': {
                    'paragraph_count': len(paragraphs),
                    'file_type': 'DOCX'
                }
            })
            
        except Exception as e:
            result['error'] = f"Error processing DOCX: {str(e)}"
            
        return result
    
    def _extract_txt(self, file_path: str = None, file_content: bytes = None) -> Dict[str, Any]:
        """Extract text from TXT files."""
        result = {'text': '', 'success': False, 'error': None, 'metadata': {}}
        
        try:
            if file_content:
                # Handle uploaded file content
                try:
                    text = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    text = file_content.decode('latin-1')
            else:
                # Handle file path
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                except UnicodeDecodeError:
                    with open(file_path, 'r', encoding='latin-1') as file:
                        text = file.read()
            
            result.update({
                'text': text,
                'success': True,
                'metadata': {
                    'character_count': len(text),
                    'file_type': 'TXT'
                }
            })
            
        except Exception as e:
            result['error'] = f"Error processing TXT: {str(e)}"
            
        return result


class TextCleaner:
    """Handles text cleaning and normalization."""
    
    def __init__(self):
        pass
    
    def clean_and_normalize(self, text: str) -> Dict[str, Any]:
        """
        Clean and normalize extracted text.
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            Dict containing cleaned text and cleaning metadata
        """
        if not text:
            return {
                'cleaned_text': '',
                'original_length': 0,
                'cleaned_length': 0,
                'cleaning_stats': {}
            }
        
        original_length = len(text)
        
        # Store original for comparison
        cleaned_text = text
        
        # Remove excessive whitespace
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        # Remove special characters but keep basic punctuation
        cleaned_text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', '', cleaned_text)
        
        # Remove multiple consecutive periods
        cleaned_text = re.sub(r'\.{3,}', '...', cleaned_text)
        
        # Clean up spacing around punctuation
        cleaned_text = re.sub(r'\s+([.,;:!?])', r'\1', cleaned_text)
        cleaned_text = re.sub(r'([.,;:!?])\s*', r'\1 ', cleaned_text)
        
        # Remove extra spaces and strip
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        
        # Calculate cleaning statistics
        cleaning_stats = {
            'whitespace_reduction': original_length - len(re.sub(r'\s+', ' ', text)),
            'special_chars_removed': len(re.findall(r'[^\w\s\.\,\;\:\!\?\-\(\)]', text)),
            'lines_count': len(text.split('\n')),
            'words_count': len(cleaned_text.split()) if cleaned_text else 0
        }
        
        return {
            'cleaned_text': cleaned_text,
            'original_length': original_length,
            'cleaned_length': len(cleaned_text),
            'cleaning_stats': cleaning_stats
        }
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract common resume/job description sections.
        
        Args:
            text (str): Cleaned text content
            
        Returns:
            Dict with identified sections
        """
        sections = {
            'contact': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'other': ''
        }
        
        # Convert to lowercase for pattern matching
        text_lower = text.lower()
        
        # Define section patterns
        patterns = {
            'contact': r'(email|phone|address|linkedin|github)',
            'summary': r'(summary|objective|profile|about)',
            'experience': r'(experience|employment|work history|professional)',
            'education': r'(education|academic|degree|university|college)',
            'skills': r'(skills|competencies|technologies|tools)'
        }
        
        # Simple section extraction based on keywords
        lines = text.split('\n')
        current_section = 'other'
        
        for line in lines:
            line_lower = line.lower().strip()
            if not line_lower:
                continue
                
            # Check if line indicates a new section
            for section, pattern in patterns.items():
                if re.search(pattern, line_lower):
                    current_section = section
                    break
            
            # Add line to current section
            if sections[current_section]:
                sections[current_section] += '\n' + line
            else:
                sections[current_section] = line
        
        return sections


# Initialize global instances
document_processor = DocumentProcessor()
text_cleaner = TextCleaner()