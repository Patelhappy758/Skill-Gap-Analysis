# file_reader_doc.py - Enhanced DOCX File Reader
"""
Enhanced DOCX File Reader with comprehensive content extraction
"""

import docx
import streamlit as st
from typing import Optional, Dict, Any, List
import os
from docx.document import Document
from docx.table import Table


def extract_text_from_paragraphs(doc: Document) -> str:
    """
    Extract text from document paragraphs
    Args:
        doc: Document object
    Returns:
        str: Extracted paragraph text
    """
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    return "\n".join(paragraphs)


def extract_text_from_tables(doc: Document) -> str:
    """
    Extract text from document tables
    Args:
        doc: Document object
    Returns:
        str: Extracted table text
    """
    tables_text = []
    
    for table_num, table in enumerate(doc.tables):
        table_content = []
        table_content.append(f"\n--- Table {table_num + 1} ---")
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            
            if row_cells:
                table_content.append(" | ".join(row_cells))
        
        if len(table_content) > 1:  # More than just the header
            tables_text.append("\n".join(table_content))
    
    return "\n".join(tables_text)


def read_docx(file_path: str, include_tables: bool = True) -> str:
    """
    Read DOCX file with comprehensive content extraction
    Args:
        file_path (str): Path to the DOCX file
        include_tables (bool): Whether to include table content
    Returns:
        str: Extracted text content from DOCX file
    """
    try:
        # Validate file
        if not file_path or not os.path.exists(file_path):
            st.error(f"❌ DOCX file not found: {file_path}")
            return ""
        
        # Get file info
        file_size = os.path.getsize(file_path)
        st.info(f"📄 Processing DOCX file: {os.path.basename(file_path)} ({file_size / 1024:.1f} KB)")
        
        # Load document
        doc = docx.Document(file_path)
        
        # Extract paragraph text
        paragraph_text = extract_text_from_paragraphs(doc)
        
        # Extract table text if requested
        table_text = ""
        if include_tables:
            table_text = extract_text_from_tables(doc)
        
        # Combine all text
        all_text_parts = [paragraph_text]
        if table_text:
            all_text_parts.append(table_text)
        
        combined_text = "\n\n".join([part for part in all_text_parts if part.strip()])
        
        # Validate extracted text
        if not combined_text.strip():
            st.warning("⚠️ No text could be extracted from the DOCX file")
            return ""
        
        # Provide extraction summary
        num_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
        num_tables = len(doc.tables)
        
        st.success(f"✅ Successfully read DOCX file ({len(combined_text)} characters)")
        st.info(f"📊 Content summary: {num_paragraphs} paragraphs, {num_tables} tables")
        
        return combined_text

    except docx.opc.exceptions.PackageNotFoundError:
        st.error("❌ Invalid DOCX file format")
        return ""
    
    except FileNotFoundError:
        st.error("❌ DOCX file not found")
        return ""
    
    except PermissionError:
        st.error("❌ Permission denied to read DOCX file")
        return ""
    
    except Exception as e:
        st.error(f"❌ Error reading DOCX file: {str(e)}")
        return ""


def get_docx_metadata(file_path: str) -> Dict[str, Any]:
    """
    Get metadata information about the DOCX file
    Args:
        file_path (str): Path to the DOCX file
    Returns:
        dict: DOCX metadata
    """
    try:
        metadata = {}
        
        # File system metadata
        stat = os.stat(file_path)
        metadata.update({
            'file_name': os.path.basename(file_path),
            'file_size': stat.st_size,
            'created_time': stat.st_ctime,
            'modified_time': stat.st_mtime
        })
        
        # DOCX-specific metadata
        try:
            doc = docx.Document(file_path)
            
            # Document properties
            props = doc.core_properties
            metadata.update({
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'keywords': props.keywords or '',
                'created': props.created,
                'modified': props.modified,
                'num_paragraphs': len([p for p in doc.paragraphs if p.text.strip()]),
                'num_tables': len(doc.tables)
            })
        except:
            pass
        
        return metadata
    except Exception as e:
        return {'error': str(e)}


